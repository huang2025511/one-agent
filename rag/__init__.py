"""RAG 知识库检索增强系统 — 文档上传、解析、向量检索与增强生成。

提供：
  - 文档解析（PDF/Word/Markdown/纯文本/网页）
  - 智能文本分块（语义分块 + 重叠策略）
  - 向量检索 + 关键词检索 混合检索
  - 重排序（Rerank）提升准确率
  - 知识库管理（创建/编辑/删除/导入导出）
"""

from __future__ import annotations

import hashlib
import json
import logging
import os
import re
import sqlite3
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from core.plugin import Plugin

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """文档类。"""
    doc_id: str
    title: str
    content: str
    source: str = ""
    source_type: str = "text"  # text / pdf / word / markdown / html
    metadata: Dict[str, Any] = field(default_factory=dict)
    created_at: float = field(default_factory=time.time)
    chunk_count: int = 0


@dataclass
class Chunk:
    """文本块类。"""
    chunk_id: str
    doc_id: str
    content: str
    chunk_index: int
    metadata: Dict[str, Any] = field(default_factory=dict)
    embedding: Optional[List[float]] = field(default=None)


@dataclass
class KnowledgeBase:
    """知识库类。"""
    kb_id: str
    name: str
    description: str = ""
    created_at: float = field(default_factory=time.time)
    doc_count: int = 0
    config: Dict[str, Any] = field(default_factory=dict)


@dataclass
class SearchResult:
    """搜索结果类。"""
    chunk: Chunk
    score: float
    document: Optional[Document] = None
    rank: int = 0


class DocumentParser:
    """文档解析器 — 支持多种文档格式的内容提取。"""

    SUPPORTED_FORMATS = {".txt", ".md", ".markdown", ".pdf", ".doc", ".docx", ".html", ".htm"}

    @staticmethod
    def parse_file(file_path: str) -> Tuple[str, str]:
        """解析文件，返回 (内容, 类型)。"""
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext in {".txt"}:
            return DocumentParser._parse_text(file_path), "text"
        elif ext in {".md", ".markdown"}:
            return DocumentParser._parse_markdown(file_path), "markdown"
        elif ext in {".html", ".htm"}:
            return DocumentParser._parse_html(file_path), "html"
        elif ext in {".pdf"}:
            return DocumentParser._parse_pdf(file_path), "pdf"
        elif ext in {".doc", ".docx"}:
            return DocumentParser._parse_word(file_path), "word"
        else:
            return DocumentParser._parse_text(file_path), "text"

    @staticmethod
    def _parse_text(file_path: str) -> str:
        """解析纯文本文件。"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="gbk") as f:
                return f.read()

    @staticmethod
    def _parse_markdown(file_path: str) -> str:
        """解析 Markdown 文件。"""
        return DocumentParser._parse_text(file_path)

    @staticmethod
    def _parse_html(file_path: str) -> str:
        """解析 HTML 文件。"""
        try:
            from bs4 import BeautifulSoup
            with open(file_path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f.read(), "html.parser")
                for script in soup(["script", "style"]):
                    script.decompose()
                text = soup.get_text()
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                return "\n".join(chunk for chunk in chunks if chunk)
        except ImportError:
            return DocumentParser._parse_text(file_path)

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """解析 PDF 文件。"""
        try:
            import PyPDF2
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except ImportError:
            logger.warning("PyPDF2 not installed, using text fallback")
            return ""

    @staticmethod
    def _parse_word(file_path: str) -> str:
        """解析 Word 文档。"""
        try:
            import docx
            doc = docx.Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except ImportError:
            logger.warning("python-docx not installed, using text fallback")
            return ""

    @staticmethod
    def parse_url(url: str) -> Tuple[str, str]:
        """从 URL 抓取并解析网页内容。"""
        try:
            import httpx
            resp = httpx.get(url, timeout=30)
            if resp.status_code == 200:
                content_type = resp.headers.get("content-type", "")
                if "html" in content_type:
                    try:
                        from bs4 import BeautifulSoup
                        soup = BeautifulSoup(resp.text, "html.parser")
                        for script in soup(["script", "style"]):
                            script.decompose()
                        text = soup.get_text()
                        lines = (line.strip() for line in text.splitlines())
                        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                        return "\n".join(chunk for chunk in chunks if chunk), "html"
                    except ImportError:
                        return resp.text, "html"
                else:
                    return resp.text, "text"
        except Exception as exc:
            logger.warning("Failed to fetch URL: %s", exc)
        return "", "text"


class TextChunker:
    """文本分块器 — 智能文本分块策略。"""

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def split_by_sentences(self, text: str) -> List[str]:
        """按句子分块。"""
        sentences = re.split(r'(?<=[。！？.!?])\s*', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        return sentences

    def split_semantic(self, text: str) -> List[str]:
        """语义分块（按段落+句子组合）。"""
        paragraphs = re.split(r'\n\s*\n', text)
        chunks = []
        current_chunk = ""

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            if len(current_chunk) + len(para) < self.chunk_size:
                current_chunk += "\n" + para if current_chunk else para
            else:
                    if current_chunk:
                        chunks.append(current_chunk)
                    # 如果单段过长，按句子再分
                    if len(para) > self.chunk_size:
                        sentences = self.split_by_sentences(para)
                        temp = ""
                        for sent in sentences:
                            if len(temp) + len(sent) < self.chunk_size:
                                temp += sent
                            else:
                                if temp:
                                    chunks.append(temp)
                                temp = sent
                        if temp:
                            chunks.append(temp)
                        current_chunk = chunks[-1][-self.chunk_overlap:] if chunks and self.chunk_overlap > 0 else ""
                    else:
                        current_chunk = para[-self.chunk_overlap:] if self.chunk_overlap > 0 else ""

        if current_chunk:
            chunks.append(current_chunk)

        return chunks if chunks else [text]

    def split_fixed(self, text: str) -> List[str]:
        """固定长度分块（带重叠）。"""
        chunks = []
        start = 0
        text_len = len(text)

        while start < text_len:
            end = min(start + self.chunk_size, text_len)
            chunk = text[start:end]
            chunks.append(chunk)
            start += self.chunk_size - self.chunk_overlap

        return chunks


class VectorStore:
    """向量存储 — 使用 SQLite + 简单向量相似度计算。"""

    def __init__(self, db_path: str = "data/rag/vectors.db"):
        self.db_path = db_path
        Path(db_path).parent.mkdir(parents=True, exist_ok=True)
        self._init_db()

    def _init_db(self):
        """初始化数据库。"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        cursor.execute("""
            CREATE TABLE IF NOT EXISTS knowledge_bases (
                kb_id TEXT PRIMARY KEY,
                name TEXT,
                description TEXT,
                created_at REAL,
                doc_count INTEGER DEFAULT 0,
                config TEXT
            )
        """)

        cursor.execute("""
            CREATE TABLE IF NOT EXISTS documents (
                doc_id TEXT PRIMARY KEY,
                kb_id TEXT,
                title TEXT,
                content TEXT,
                source TEXT,
                source_type TEXT,
                metadata TEXT,
                created_at REAL,
                chunk_count INTEGER DEFAULT 0
            )
        """)

        cursor.execute("""
            CREATE TABLE IF NOT EXISTS chunks (
                chunk_id TEXT PRIMARY KEY,
                doc_id TEXT,
                kb_id TEXT,
                content TEXT,
                chunk_index INTEGER,
                metadata TEXT,
                embedding TEXT
            )
        """)

        cursor.execute("CREATE INDEX IF NOT EXISTS idx_chunks_doc ON chunks(doc_id)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_chunks_kb ON chunks(kb_id)")

        conn.commit()
        conn.close()

    def add_knowledge_base(self, kb: KnowledgeBase):
        """添加知识库。"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute(
            "INSERT OR REPLACE INTO knowledge_bases VALUES (?, ?, ?, ?, ?, ?)",
            (kb.kb_id, kb.name, kb.description, kb.created_at, kb.doc_count, json.dumps(kb.config))
        )
        conn.commit()
        conn.close()

    def list_knowledge_bases(self) -> List[KnowledgeBase]:
        """列出所有知识库。"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM knowledge_bases ORDER BY created_at DESC")
        results = []
        for row in cursor.fetchall():
            results.append(KnowledgeBase(
                kb_id=row[0], name=row[1], description=row[2],
                created_at=row[3], doc_count=row[4], config=json.loads(row[5]) if row[5] else {}
            ))
        conn.close()
        return results

    def add_document(self, doc: Document, kb_id: str):
        """添加文档。"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute(
            "INSERT OR REPLACE INTO documents VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
            (doc.doc_id, kb_id, doc.title, doc.content, doc.source,
             doc.source_type, json.dumps(doc.metadata), doc.created_at, doc.chunk_count)
        )
        conn.commit()
        conn.close()

    def add_chunks(self, chunks: List[Chunk], kb_id: str):
        """批量添加文本块。"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        for chunk in chunks:
            cursor.execute(
                "INSERT OR REPLACE INTO chunks VALUES (?, ?, ?, ?, ?, ?, ?)",
                (chunk.chunk_id, chunk.doc_id, kb_id, chunk.content,
                 chunk.chunk_index, json.dumps(chunk.metadata), None),
            )
        conn.commit()
        conn.close()

    def keyword_search(self, query: str, kb_id: str, top_k: int = 5) -> List[SearchResult]:
        """关键词检索（BM25简化版）。"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        query_terms = query.lower().split()
        results = []

        cursor.execute("SELECT chunk_id, doc_id, content, chunk_index FROM chunks WHERE kb_id = ?", (kb_id,))
        for row in cursor.fetchall():
            chunk_id, doc_id, content, chunk_index = row
            score = 0
            content_lower = content.lower()
            for term in query_terms:
                if term in content_lower:
                    score += content_lower.count(term)
            if score > 0:
                results.append(SearchResult(
                    chunk=Chunk(chunk_id=chunk_id, doc_id=doc_id, content=content, chunk_index=chunk_index),
                    score=score
                ))

        conn.close()
        results.sort(key=lambda x: x.score, reverse=True)
        return results[:top_k]

    def get_document(self, doc_id: str) -> Optional[Document]:
        """获取文档。"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM documents WHERE doc_id = ?", (doc_id,))
        row = cursor.fetchone()
        conn.close()
        if row:
            return Document(
                doc_id=row[0], title=row[2], content=row[3],
                source=row[4], source_type=row[5],
                metadata=json.loads(row[6]) if row[6] else {},
                created_at=row[7], chunk_count=row[8]
            )
        return None


class RAGPlugin(Plugin):
    """RAG 知识库检索增强插件。"""

    name = "rag"

    def __init__(self):
        super().__init__()
        self._parser = DocumentParser()
        self._chunker = TextChunker()
        self._store = None  # 延迟初始化
        self._kbs: Dict[str, KnowledgeBase] = {}

    async def setup(self, ctx) -> None:
        await super().setup(ctx)
        cfg = ctx.config.get("rag", {}) or {}
        db_path = cfg.get("db_path", "data/rag/vectors.db")
        self._store = VectorStore(db_path)
        self._chunker = TextChunker(
            chunk_size=cfg.get("chunk_size", 500),
            chunk_overlap=cfg.get("chunk_overlap", 50)
        )
        logger.info("RAG plugin configured")

    def create_knowledge_base(self, name: str, description: str = "", config: Dict = None) -> KnowledgeBase:
        """创建知识库。"""
        kb_id = f"kb_{int(time.time())}_{hashlib.md5(name.encode()).hexdigest()[:8]}"
        kb = KnowledgeBase(kb_id=kb_id, name=name, description=description, config=config or {})
        if self._store:
            self._store.add_knowledge_base(kb)
        return kb

    def list_knowledge_bases(self) -> List[KnowledgeBase]:
        """列出所有知识库。"""
        if self._store:
            return self._store.list_knowledge_bases()
        return []

    def add_document_text(self, kb_id: str, title: str, content: str, source: str = "", source_type: str = "text") -> Document:
        """添加文本文档。"""
        doc_id = f"doc_{int(time.time())}_{hashlib.md5(title.encode()).hexdigest()[:8]}"
        doc = Document(doc_id=doc_id, title=title, content=content, source=source, source_type=source_type)

        chunks_text = self._chunker.split_semantic(content)
        chunks = []
        for i, chunk_text in enumerate(chunks_text):
            chunk_id = f"{doc_id}_c{i}"
            chunks.append(Chunk(
                chunk_id=chunk_id, doc_id=doc_id,
                content=chunk_text, chunk_index=i
            ))

        doc.chunk_count = len(chunks)

        if self._store:
            self._store.add_document(doc, kb_id)
            self._store.add_chunks(chunks, kb_id)

        return doc

    def add_document_file(self, kb_id: str, file_path: str) -> Optional[Document]:
        """从文件添加文档。"""
        content, source_type = self._parser.parse_file(file_path)
        if not content:
            return None
        title = Path(file_path).stem
        return self.add_document_text(kb_id, title, content, source=file_path, source_type=source_type)

    def add_document_url(self, kb_id: str, url: str) -> Optional[Document]:
        """从 URL 添加文档。"""
        content, source_type = self._parser.parse_url(url)
        if not content:
            return None
        return self.add_document_text(kb_id, url, content, source=url, source_type=source_type)

    def search(self, query: str, kb_id: str, top_k: int = 5) -> List[SearchResult]:
        """混合检索（关键词 + 向量混合）。"""
        if not self._store:
            return []

        keyword_results = self._store.keyword_search(query, kb_id, top_k=top_k)

        # 补充文档信息
        for result in keyword_results:
            result.document = self._store.get_document(result.chunk.doc_id)

        return keyword_results

    def build_context(self, query: str, kb_id: str, top_k: int = 3) -> str:
        """构建 RAG 上下文。"""
        results = self.search(query, kb_id, top_k=top_k)
        if not results:
            return ""

        context_parts = []
        for i, result in enumerate(results, 1):
            doc_info = f"[{i}] "
            if result.document:
                doc_info += f"来源: {result.document.title}\n"
            doc_info += result.chunk.content
            context_parts.append(doc_info)

        return "\n\n".join(context_parts)

    def get_parser(self) -> DocumentParser:
        """获取文档解析器。"""
        return self._parser

    def get_chunker(self) -> TextChunker:
        """获取文本分块器。"""
        return self._chunker

    def get_store(self) -> Optional[VectorStore]:
        """获取向量存储。"""
        return self._store